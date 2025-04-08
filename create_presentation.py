from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_number(doc, text, level):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{text}")
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14)
    return paragraph

def add_subheading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return paragraph

def add_normal_text(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    return paragraph

def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return paragraph

def create_presentation():
    doc = Document()
    
    # Title
    title = doc.add_heading('Hệ Thống Đặt Vé Xem Phim', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    add_heading_with_number(doc, "1. Giới thiệu", 1)
    add_normal_text(doc, "Hệ thống đặt vé xem phim là một ứng dụng web được phát triển bằng Spring Boot, cung cấp các chức năng cần thiết cho việc đặt vé xem phim trực tuyến. Hệ thống được thiết kế với kiến trúc MVC, sử dụng các công nghệ hiện đại và tuân thủ các best practices trong phát triển phần mềm.")
    
    # Features
    add_heading_with_number(doc, "2. Các chức năng chính", 1)
    
    # Authentication
    add_heading_with_number(doc, "2.1. Xác thực người dùng", 2)
    add_normal_text(doc, "Hệ thống cung cấp các chức năng xác thực người dùng đầy đủ:")
    add_normal_text(doc, "• Đăng ký tài khoản mới với xác thực email")
    add_normal_text(doc, "• Đăng nhập với tài khoản đã đăng ký")
    add_normal_text(doc, "• Đăng nhập bằng Google và Facebook")
    add_normal_text(doc, "• Quản lý phiên đăng nhập và bảo mật")
    
    # User Profile
    add_heading_with_number(doc, "2.2. Quản lý thông tin cá nhân", 2)
    add_normal_text(doc, "Người dùng có thể quản lý thông tin cá nhân của mình:")
    add_normal_text(doc, "• Cập nhật thông tin cá nhân (họ tên, số điện thoại, địa chỉ)")
    add_normal_text(doc, "• Đổi mật khẩu với xác thực mật khẩu cũ")
    add_normal_text(doc, "• Xác thực thẻ sinh viên để hưởng ưu đãi")
    
    # Movie Management
    add_heading_with_number(doc, "2.3. Quản lý phim", 2)
    add_normal_text(doc, "Hệ thống cho phép quản lý thông tin phim:")
    add_normal_text(doc, "• Hiển thị danh sách phim đang chiếu")
    add_normal_text(doc, "• Xem chi tiết thông tin phim")
    add_normal_text(doc, "• Phân loại phim theo thể loại")
    add_normal_text(doc, "• Tìm kiếm phim theo tên hoặc thể loại")
    
    # Booking System
    add_heading_with_number(doc, "2.4. Hệ thống đặt vé", 2)
    add_normal_text(doc, "Chức năng đặt vé được thiết kế với nhiều tính năng:")
    add_normal_text(doc, "• Chọn suất chiếu phim")
    add_normal_text(doc, "• Chọn ghế ngồi trực quan")
    add_normal_text(doc, "• Phân biệt vé người lớn và vé sinh viên")
    add_normal_text(doc, "• Tính toán giá vé tự động")
    add_normal_text(doc, "• Thanh toán trực tuyến qua VNPAY")
    
    # Technical Details
    add_heading_with_number(doc, "3. Chi tiết kỹ thuật", 1)
    
    # Architecture
    add_heading_with_number(doc, "3.1. Kiến trúc hệ thống", 2)
    add_normal_text(doc, "Hệ thống được xây dựng theo mô hình MVC (Model-View-Controller):")
    add_normal_text(doc, "• Model: Các entity và repository")
    add_normal_text(doc, "• View: Các template HTML sử dụng Thymeleaf")
    add_normal_text(doc, "• Controller: Xử lý request và điều hướng")
    
    # Security
    add_heading_with_number(doc, "3.2. Bảo mật", 2)
    add_normal_text(doc, "Các biện pháp bảo mật được triển khai:")
    add_normal_text(doc, "• Mã hóa mật khẩu với BCrypt")
    add_normal_text(doc, "• Xác thực JWT cho API")
    add_normal_text(doc, "• Phân quyền người dùng với Spring Security")
    add_normal_text(doc, "• Bảo vệ CSRF")
    
    # Testing
    add_heading_with_number(doc, "3.3. Kiểm thử", 2)
    add_normal_text(doc, "Hệ thống được kiểm thử kỹ lưỡng:")
    add_normal_text(doc, "• Unit test cho các controller và service")
    add_normal_text(doc, "• Integration test cho các API")
    add_normal_text(doc, "• Test coverage cao")
    
    # FAQ
    add_heading_with_number(doc, "4. Câu hỏi thường gặp", 1)
    
    # FAQ Items
    faq_items = [
        ("Làm sao để đặt vé sinh viên?", 
         "Để đặt vé sinh viên, bạn cần: 1) Đăng ký tài khoản với email sinh viên, 2) Cập nhật ảnh thẻ sinh viên trong phần profile, 3) Chọn loại vé sinh viên khi đặt vé."),
        
        ("Cách thanh toán vé?", 
         "Hệ thống hỗ trợ thanh toán trực tuyến qua VNPAY. Sau khi chọn ghế, bạn sẽ được chuyển đến trang thanh toán VNPAY để hoàn tất giao dịch."),
        
        ("Làm sao để thay đổi thông tin cá nhân?", 
         "Bạn có thể thay đổi thông tin cá nhân trong phần Profile. Để đổi mật khẩu, bạn cần nhập mật khẩu cũ để xác thực."),
        
        ("Có thể hủy vé đã đặt không?", 
         "Vé đã đặt và thanh toán không thể hủy. Tuy nhiên, bạn có thể liên hệ với chúng tôi trong trường hợp đặc biệt."),
        
        ("Làm sao để xem lịch sử đặt vé?", 
         "Bạn có thể xem lịch sử đặt vé trong phần Profile của mình. Tại đây, bạn sẽ thấy danh sách các vé đã đặt và trạng thái của chúng.")
    ]
    
    for question, answer in faq_items:
        add_subheading(doc, question)
        add_normal_text(doc, answer)
    
    # Save the document
    doc.save('Movie_Ticket_System_Presentation.docx')

if __name__ == '__main__':
    create_presentation() 